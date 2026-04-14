"""Customer requirements file parser.

Reads requirements from CSV, Excel, Markdown, or Word files and returns
structured CustomerRequirement objects. Designed to be forgiving: if it
can extract at least title and description, that's enough.
"""

import csv
import io
import re
from pathlib import Path

from crewai.tools import BaseTool
from pydantic import BaseModel, Field


# ---------- Data model ----------


class CustomerRequirement(BaseModel):
    id: str = ""
    title: str
    description: str
    priority: str = "unspecified"
    category: str = ""
    notes: str = ""
    source: str = "customer_input"


# ---------- Priority mapping ----------

_PRIORITY_MAP = {
    "p0": "P0", "0": "P0", "must have": "P0", "must-have": "P0",
    "critical": "P0", "highest": "P0",
    "p1": "P1", "1": "P1", "should have": "P1", "should-have": "P1",
    "high": "P1",
    "p2": "P2", "2": "P2", "could have": "P2", "could-have": "P2",
    "medium": "P2", "nice to have": "P2", "nice-to-have": "P2",
    "p3": "P3", "3": "P3", "low": "P3", "future": "P3",
    "future consideration": "P3",
    "p4": "P4", "4": "P4", "backlog": "P4", "lowest": "P4",
    "won't have": "P4", "wont have": "P4",
}


def _normalize_priority(raw: str) -> str:
    if not raw or not raw.strip():
        return "unspecified"
    key = raw.strip().lower()
    return _PRIORITY_MAP.get(key, raw.strip())


# ---------- Column matching ----------

_TITLE_PATTERNS = ["title", "name", "story", "summary", "requirement name", "user story"]
_DESC_PATTERNS = ["description", "detail", "requirement", "body", "text", "acceptance"]
_PRIORITY_PATTERNS = ["priority", "importance", "severity", "p0", "p1", "p2", "p3", "p4", "level"]
_ID_PATTERNS = ["id", "key", "number", "ref", "ticket"]
_CATEGORY_PATTERNS = ["category", "group", "theme", "type", "epic", "module"]
_NOTES_PATTERNS = ["notes", "comment", "remark", "additional"]


def _match_column(header: str, patterns: list[str]) -> bool:
    h = header.strip().lower()
    return any(p in h for p in patterns)


def _map_columns(headers: list[str]) -> dict[str, int | None]:
    mapping: dict[str, int | None] = {
        "title": None, "description": None, "priority": None,
        "id": None, "category": None, "notes": None,
    }
    for i, h in enumerate(headers):
        if mapping["title"] is None and _match_column(h, _TITLE_PATTERNS):
            mapping["title"] = i
        elif mapping["description"] is None and _match_column(h, _DESC_PATTERNS):
            mapping["description"] = i
        elif mapping["priority"] is None and _match_column(h, _PRIORITY_PATTERNS):
            mapping["priority"] = i
        elif mapping["id"] is None and _match_column(h, _ID_PATTERNS):
            mapping["id"] = i
        elif mapping["category"] is None and _match_column(h, _CATEGORY_PATTERNS):
            mapping["category"] = i
        elif mapping["notes"] is None and _match_column(h, _NOTES_PATTERNS):
            mapping["notes"] = i
    return mapping


def _row_to_requirement(row: list[str], col_map: dict[str, int | None], index: int) -> CustomerRequirement | None:
    def _get(field: str) -> str:
        idx = col_map.get(field)
        if idx is not None and idx < len(row):
            return str(row[idx]).strip()
        return ""

    title = _get("title")
    desc = _get("description")
    if not title and not desc:
        return None

    return CustomerRequirement(
        id=_get("id") or f"CR-{index + 1:03d}",
        title=title or desc[:80],
        description=desc or title,
        priority=_normalize_priority(_get("priority")),
        category=_get("category"),
        notes=_get("notes"),
    )


# ---------- Format parsers ----------


def _parse_csv(file_path: Path) -> list[CustomerRequirement]:
    text = file_path.read_text(encoding="utf-8")
    reader = csv.reader(io.StringIO(text))
    rows = list(reader)
    if len(rows) < 2:
        return []

    col_map = _map_columns(rows[0])
    if col_map["title"] is None and col_map["description"] is None:
        return []

    results = []
    for i, row in enumerate(rows[1:]):
        req = _row_to_requirement(row, col_map, i)
        if req:
            results.append(req)
    return results


def _parse_excel(file_path: Path) -> list[CustomerRequirement]:
    import openpyxl

    wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
    ws = wb.active
    if ws is None:
        return []

    rows = []
    for row in ws.iter_rows(values_only=True):
        rows.append([str(c) if c is not None else "" for c in row])
    wb.close()

    if len(rows) < 2:
        return []

    col_map = _map_columns(rows[0])
    if col_map["title"] is None and col_map["description"] is None:
        return []

    results = []
    for i, row in enumerate(rows[1:]):
        req = _row_to_requirement(row, col_map, i)
        if req:
            results.append(req)
    return results


def _parse_markdown(file_path: Path) -> list[CustomerRequirement]:
    text = file_path.read_text(encoding="utf-8")
    return _parse_markdown_text(text)


def _parse_markdown_text(text: str) -> list[CustomerRequirement]:
    # Try table first
    table_reqs = _parse_markdown_table(text)
    if table_reqs:
        return table_reqs

    # Fall back to header-based parsing
    return _parse_markdown_headers(text)


def _parse_markdown_table(text: str) -> list[CustomerRequirement]:
    lines = text.strip().split("\n")
    table_start = None
    for i, line in enumerate(lines):
        if "|" in line and i + 1 < len(lines) and re.match(r"^\s*\|[\s\-:|]+\|", lines[i + 1]):
            table_start = i
            break

    if table_start is None:
        return []

    header_line = lines[table_start]
    headers = [h.strip() for h in header_line.split("|") if h.strip()]
    col_map = _map_columns(headers)
    if col_map["title"] is None and col_map["description"] is None:
        return []

    results = []
    for i, line in enumerate(lines[table_start + 2:]):
        if "|" not in line:
            break
        cells = [c.strip() for c in line.split("|") if c.strip() or line.count("|") > 1]
        # Handle leading/trailing pipes
        raw_cells = line.split("|")
        cells = [c.strip() for c in raw_cells[1:-1]] if raw_cells[0].strip() == "" else [c.strip() for c in raw_cells]
        req = _row_to_requirement(cells, col_map, i)
        if req:
            results.append(req)
    return results


_PRIORITY_RE = re.compile(r"\b(P[0-4])\b|Priority:\s*(\w+)|(\b(?:Must|Should|Could|Won'?t)\s+have\b)", re.IGNORECASE)


def _extract_priority_from_text(text: str) -> str:
    m = _PRIORITY_RE.search(text)
    if m:
        val = m.group(1) or m.group(2) or m.group(3)
        return _normalize_priority(val)
    return "unspecified"


def _parse_markdown_headers(text: str) -> list[CustomerRequirement]:
    lines = text.split("\n")
    results = []
    current_title = ""
    current_body: list[str] = []
    current_category = ""
    header_level = 0

    def _flush():
        if current_title and current_body:
            body = "\n".join(current_body).strip()
            # Check for list items under a header
            list_items = [l.strip().lstrip("-*").strip() for l in current_body if re.match(r"^\s*[-*]\s+", l)]
            if list_items and len(list_items) > 1:
                for j, item in enumerate(list_items):
                    results.append(CustomerRequirement(
                        id=f"CR-{len(results) + 1:03d}",
                        title=item[:80],
                        description=item,
                        priority=_extract_priority_from_text(item),
                        category=current_title,
                    ))
            else:
                results.append(CustomerRequirement(
                    id=f"CR-{len(results) + 1:03d}",
                    title=current_title,
                    description=body,
                    priority=_extract_priority_from_text(body),
                    category=current_category,
                ))

    for line in lines:
        header_match = re.match(r"^(#{2,3})\s+(.+)", line)
        if header_match:
            _flush()
            level = len(header_match.group(1))
            title = header_match.group(2).strip()
            if level == 2:
                current_category = title
                current_title = title
            else:
                current_title = title
            current_body = []
            header_level = level
        elif line.strip():
            current_body.append(line)

    _flush()
    return results


def _parse_docx(file_path: Path) -> list[CustomerRequirement]:
    import docx

    doc = docx.Document(str(file_path))

    # Try tables first
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        col_map = _map_columns(headers)
        if col_map["title"] is None and col_map["description"] is None:
            continue

        results = []
        for i, row in enumerate(table.rows[1:]):
            cells = [cell.text.strip() for cell in row.cells]
            req = _row_to_requirement(cells, col_map, i)
            if req:
                results.append(req)
        if results:
            return results

    # Fall back to text-based parsing using markdown heuristics
    text_lines = []
    for para in doc.paragraphs:
        style_name = (para.style.name or "").lower() if para.style else ""
        if "heading 2" in style_name or "heading 3" in style_name:
            level = "##" if "heading 2" in style_name else "###"
            text_lines.append(f"{level} {para.text}")
        elif para.text.strip():
            text_lines.append(para.text)

    if text_lines:
        return _parse_markdown_text("\n".join(text_lines))

    return []


# ---------- Main entry point ----------

_PARSERS = {
    ".csv": _parse_csv,
    ".xlsx": _parse_excel,
    ".xls": _parse_excel,
    ".md": _parse_markdown,
    ".docx": _parse_docx,
}


def parse_requirements_file(file_path: str) -> list[CustomerRequirement]:
    """Parse a requirements file and return structured requirements.

    Supports: .csv, .xlsx, .xls, .md, .docx
    Required elements: title, description, priority (priority can default to 'unspecified')
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"Requirements file not found: {file_path}")

    ext = path.suffix.lower()
    parser = _PARSERS.get(ext)
    if parser is None:
        raise ValueError(
            f"Could not parse requirements from {path.name}. "
            f"Expected a file containing requirement titles, descriptions, and priorities "
            f"in CSV, Excel, Markdown, or Word format."
        )

    try:
        results = parser(path)
    except Exception as e:
        raise ValueError(
            f"Could not parse requirements from {path.name}. "
            f"Expected a file containing requirement titles, descriptions, and priorities "
            f"in CSV, Excel, Markdown, or Word format. Error: {e}"
        ) from e

    if not results:
        raise ValueError(
            f"Could not parse requirements from {path.name}. "
            f"Expected a file containing requirement titles, descriptions, and priorities "
            f"in CSV, Excel, Markdown, or Word format."
        )

    return results


# ---------- CrewAI tool wrapper ----------


class RequirementsReaderTool(BaseTool):
    name: str = "requirements_reader"
    description: str = (
        "Reads a customer requirements file (CSV, Excel, Markdown, or Word) "
        "and returns structured requirements with titles, descriptions, and priorities."
    )

    def _run(self, file_path: str) -> str:
        requirements = parse_requirements_file(file_path)
        blocks = []
        for req in requirements:
            lines = [
                f"[{req.id}] Title: {req.title}",
                f"Description: {req.description}",
                f"Priority: {req.priority}",
            ]
            if req.category:
                lines.append(f"Category: {req.category}")
            if req.notes:
                lines.append(f"Notes: {req.notes}")
            blocks.append("\n".join(lines))
        return "\n---\n".join(blocks)
